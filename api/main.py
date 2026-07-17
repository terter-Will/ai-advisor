from fastapi import FastAPI, HTTPException, Query, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import pymysql, os, json, uuid, io
import httpx
from typing import Optional, Literal

# ==================== env & CORS ====================
load_dotenv(encoding='utf-8-sig')
allowed = os.getenv("ALLOW_ORIGINS", "http://localhost:5173, http://127.0.0.1:5173")
allow_list = [o.strip() for o in allowed.split(",") if o.strip()]

OPENWEBUI_BASE_URL = os.getenv("OPENWEBUI_BASE_URL", "http://172.16.188.175:3000")
OPENWEBUI_API_KEY  = os.getenv("OPENWEBUI_API_KEY", "")
OPENWEBUI_MODEL    = os.getenv("OPENWEBUI_MODEL", "nttdata-taiwan---ai-advisor")
OPENWEBUI_KB_ID    = os.getenv("OPENWEBUI_KB_ID", "311af74a-1911-484f-8b47-b13eaabd89a9")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== DB ====================
def get_conn():
    return pymysql.connect(
        host=os.getenv("DB_HOST", "127.0.0.1"),
        port=int(os.getenv("DB_PORT", "3307")),
        user=os.getenv("DB_USER", "admin"),
        password=os.getenv("DB_PASS", "1qaz@WSX"),
        database=os.getenv("DB_NAME", "ai_advisor"),
        cursorclass=pymysql.cursors.DictCursor,
        autocommit=True,
        charset="utf8mb4",
    )

# ==================== Schemas ====================
class LoginIn(BaseModel):
    userid: str
    password: str

class UserCreate(BaseModel):
    userid: str
    password: str
    role: str = "user"
    # SAP
    sapconn: Optional[str] = None
    sapaddr: Optional[str] = None
    sapins:  Optional[str] = None
    sapid:   Optional[str] = None
    sapclnt: Optional[str] = None
    saprouter: Optional[str] = None
    sapusr:    Optional[str] = None
    sappw:     Optional[str] = None
    # Points
    points_balance: Optional[int] = None
    is_active: Optional[int] = 1

class UserUpdate(BaseModel):
    password: Optional[str] = None
    role: Optional[str] = None
    sapconn: Optional[str] = None
    sapaddr: Optional[str] = None
    sapins:  Optional[str] = None
    sapid:   Optional[str] = None
    sapclnt: Optional[str] = None
    saprouter: Optional[str] = None
    sapusr:    Optional[str] = None
    sappw:     Optional[str] = None
    points_balance: Optional[int] = None
    is_active: Optional[int] = None

def row_to_public(row: dict) -> dict:
    return {
        "userid":   row["userid"],
        "role":     row["role"],
        "sapconn":  row["sapconn"],
        "sapaddr":  row["sapaddr"],
        "sapins":   row["sapins"],
        "sapid":    row["sapid"],
        "sapclnt":  row["sapclnt"],
        "saprouter":row["saprouter"],
        "sapusr":   row["sapusr"],
        "points_balance": row["points_balance"],
        "is_active":      row["is_active"],
        "created_at":     row["created_at"],
    }

# ==================== Auth ====================
@app.post("/api/login")
def login(payload: LoginIn):
    sql = """
      SELECT id, userid, password, role, points_balance, is_active
      FROM USRINFO WHERE userid=%s LIMIT 1
    """
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(sql, (payload.userid,))
        row = cur.fetchone()

    if not row or row["password"] != payload.password:
        raise HTTPException(status_code=401, detail="invalid credentials")
    if not row["is_active"]:
        raise HTTPException(status_code=403, detail="user disabled")

    return {"user": {
        "id": row["id"],
        "userid": row["userid"],
        "role": row["role"],
        "points_balance": row["points_balance"],
    }}

# ==================== Users CRUD ====================
@app.get("/api/users")
def list_users(q: Optional[str] = None, limit: int = 50):
    """依 userid 模糊查詢（做搜尋/自動完成用）"""
    limit = max(1, min(100, limit))
    with get_conn() as conn, conn.cursor() as cur:
        if q:
            cur.execute("""SELECT userid, role, points_balance
                           FROM USRINFO
                           WHERE userid LIKE %s
                           ORDER BY userid
                           LIMIT %s""", (f"%{q}%", limit))
        else:
            cur.execute("""SELECT userid, role, points_balance
                           FROM USRINFO
                           ORDER BY userid
                           LIMIT %s""", (limit,))
        return cur.fetchall()

@app.get("/api/users/{userid}")
def get_user(userid: str):
    sql = """
      SELECT userid, role,
             sapconn, sapaddr, sapins, sapid, sapclnt, saprouter, sapusr,
             password, sappw,
             points_balance, points_updated_at, is_active, created_at
      FROM USRINFO WHERE userid=%s LIMIT 1
    """
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(sql, (userid,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="not found")
        return row

@app.post("/api/users")
def create_user(body: UserCreate):
    sql = """
      INSERT INTO USRINFO
        (userid, password, role,
         sapconn, sapaddr, sapins, sapid, sapclnt, saprouter, sapusr, sappw,
         points_balance, is_active)
      VALUES (%s,%s,%s, %s,%s,%s,%s,%s,%s,%s,%s, %s,%s)
    """
    params = (
        body.userid, body.password, body.role,
        body.sapconn, body.sapaddr, body.sapins, body.sapid, body.sapclnt,
        body.saprouter, body.sapusr, body.sappw,
        body.points_balance if body.points_balance is not None else 0,
        body.is_active if body.is_active is not None else 1
    )
    try:
        with get_conn() as conn, conn.cursor() as cur:
            cur.execute(sql, params)
    except pymysql.err.IntegrityError:
        raise HTTPException(status_code=409, detail="userid exists")
    return get_user(body.userid)

@app.put("/api/users/{userid}")
def update_user(userid: str, body: UserUpdate):
    fields, params = [], []
    for col in ["password","role","sapconn","sapaddr","sapins","sapid",
                "sapclnt","saprouter","sapusr","sappw","points_balance","is_active"]:
        val = getattr(body, col)
        if val is not None and not (isinstance(val, str) and val == ""):
            fields.append(f"{col}=%s"); params.append(val)
    if not fields:
        return get_user(userid)

    sql = f"""
      UPDATE USRINFO
         SET {', '.join(fields)},
             points_updated_at = IFNULL(points_updated_at, NOW())
       WHERE userid=%s LIMIT 1
    """
    params.append(userid)
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(sql, params)
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="not found")
    return get_user(userid)

@app.delete("/api/users/{userid}")
def delete_user(userid: str):
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("DELETE FROM USRINFO WHERE userid=%s LIMIT 1", (userid,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="not found")
    return {"ok": True}

# ==================== SAP RFC READ TABLE（保留） ====================
REQUIRED_SAP_FIELDS = ["ashost","sysnr","client","user","passwd"]

def _fetch_user_sap_params(userid: str) -> dict:
    sql = """
      SELECT
        sapaddr   AS ashost,
        sapins    AS sysnr,
        sapclnt   AS client,
        sapusr    AS "user",
        sappw     AS passwd,
        saprouter AS saprouter
      FROM USRINFO WHERE userid=%s LIMIT 1
    """
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(sql, (userid,))
        row = cur.fetchone()
    if not row:
        raise HTTPException(404, "查無此使用者")

    for k, v in list(row.items()):
        if isinstance(v, str): row[k] = v.strip()

    missing = [k for k in REQUIRED_SAP_FIELDS if not row.get(k)]
    if missing:
        raise HTTPException(400, f"此使用者尚未設定完整 SAP 連線資訊：缺少 {', '.join(missing)}")

    if not row["sysnr"].isdigit() or not row["client"].isdigit():
        raise HTTPException(400, "sysnr/client 應為數字")

    row["sysnr"]  = row["sysnr"].zfill(2)
    row["client"] = row["client"].zfill(3)
    if not row.get("saprouter"): row.pop("saprouter", None)
    return row

def _parse_rfc_read_table(result: dict, delimiter: str = "|") -> dict:
    fields = result.get("FIELDS", [])
    data   = result.get("DATA", [])
    columns = [f["FIELDNAME"] for f in fields]

    rows = []
    for d in data:
        wa = d.get("WA", "")
        parts = wa.split(delimiter)
        if len(parts) != len(columns):
            parts = []
            for f in fields:
                start  = int(f["OFFSET"])
                length = int(f["LENGTH"])
                parts.append(wa[start:start+length].rstrip())
        else:
            parts = [p.rstrip() for p in parts]
        rows.append(dict(zip(columns, parts)))
    return {"columns": columns, "rows": rows, "raw": {"FIELDS": fields, "DATA": data}}

def _rfc_read_table(userid: str, table: str, field_names: Optional[list] = None, rows: int = 100) -> dict:
    """內部共用：呼叫 RFC_READ_TABLE，回傳 {columns, rows, raw}（供 sap-read 與資安掃描共用）。"""
    try:
        from pyrfc import Connection, ABAPApplicationError, ABAPRuntimeError, CommunicationError, LogonError
    except Exception as e:
        raise HTTPException(500, f"pyrfc 未就緒或 SDK 未設定：{e}")

    params = _fetch_user_sap_params(userid)
    try:
        conn = Connection(**params)
    except Exception as e:
        raise HTTPException(502, f"SAP 連線失敗：{e}")

    try:
        sel_fields = [{"FIELDNAME": n.strip().upper()} for n in (field_names or []) if n.strip()]
        result = conn.call(
            "RFC_READ_TABLE",
            QUERY_TABLE=table.upper(),
            DELIMITER="|",
            ROWCOUNT=rows,
            ROWSKIPS=0,
            OPTIONS=[],
            FIELDS=sel_fields
        )
    except (ABAPApplicationError, ABAPRuntimeError, CommunicationError, LogonError) as e:
        raise HTTPException(502, f"RFC_READ_TABLE 失敗：{e}")
    finally:
        try: conn.close()
        except: pass

    return _parse_rfc_read_table(result, delimiter="|")

@app.get("/api/users/{userid}/sap-read")
def sap_read(userid: str, table: str, rows: int = 100, fields: Optional[str] = None):
    field_names = fields.split(",") if fields else []
    parsed = _rfc_read_table(userid, table, field_names, rows)
    return {"table": table.upper(), "rowcount": len(parsed["rows"]), **parsed}

@app.get("/api/users/{userid}/sap-ping")
def sap_ping(userid: str):
    try:
        from pyrfc import Connection
    except Exception as e:
        raise HTTPException(500, f"pyrfc 未就緒或 SDK 未設定：{e}")

    params = _fetch_user_sap_params(userid)
    try:
        conn = Connection(**params)
        conn.call("RFC_PING")
        attrs = conn.get_connection_attributes()
        conn.close()
        return {"ok": True, "attrs": attrs}
    except Exception as e:
        raise HTTPException(502, f"SAP 連線失敗：{e}")

# ==================== 資安漏洞檢測（RFC 讀取 + LLM 分析） ====================
# 流程：RFC 讀 PRDVERS/CVERS → 寫入 ZPRDVERS/ZCVERS（同一 batch_id）→ 呼叫 Open WebUI（挂載
# SAP Base Knowledge，透過 files 參數強制檢索，不依賴使用者手動輸入 #）→ 存 security_scans → 回傳結構化結果。
# 使用者全程只看到 AI Advisor 自己的畫面，不會接觸到 Open WebUI 的對話介面。

PRDVERS_FIELDS = ["ID", "NAME", "VERSION", "VENDOR", "DESCRIPT", "INSTSTATUS", "MOD_DATE", "MOD_TIME"]
CVERS_FIELDS   = ["COMPONENT", "RELEASE", "EXTRELEASE", "COMP_TYPE"]

SECURITY_SCAN_SYSTEM_PROMPT = """你是 SAP Basis 資安顧問。根據使用者提供的 SAP 系統版本資訊
（產品版本 PRDVERS、軟體元件版本 CVERS），判斷是否有已知的資安疑慮
（例如版本過舊、已停止維護、已知需要修補的元件等）。
請務必只輸出一個 JSON 物件，不要有其他文字、不要用 markdown code block 包起來，格式如下：
{"summary": "一到兩句話的整體摘要", "findings": [{"severity": "high|medium|low", "title": "簡短標題", "detail": "說明與建議"}]}
如果沒有發現任何疑慮，findings 請回傳空陣列，summary 說明目前版本狀況正常。"""

def _sap_date(s: Optional[str]):
    s = (s or "").strip()
    if len(s) == 8 and s.isdigit() and s != "00000000":
        return f"{s[0:4]}-{s[4:6]}-{s[6:8]}"
    return None

def _sap_time(s: Optional[str]):
    s = (s or "").strip()
    if len(s) == 6 and s.isdigit():
        return f"{s[0:2]}:{s[2:4]}:{s[4:6]}"
    return None

def _insert_prdvers_batch(cur, owner_user_id: int, batch_id: str, rows: list):
    sql = """
      INSERT INTO ZPRDVERS
        (owner_user_id, batch_id, BORM_ID, BORM_NAME, BORM_VERS, BORM_VEND, BORM_NAME1, INSTSTATE, MOD_DATE, MOD_TIME)
      VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
    """
    for r in rows:
        cur.execute(sql, (
            owner_user_id, batch_id,
            r.get("ID") or "", r.get("NAME") or "", r.get("VERSION") or "", r.get("VENDOR") or "",
            r.get("DESCRIPT") or "", (r.get("INSTSTATUS") or "")[:1] or None,
            _sap_date(r.get("MOD_DATE")), _sap_time(r.get("MOD_TIME")),
        ))

def _insert_cvers_batch(cur, owner_user_id: int, batch_id: str, rows: list):
    sql = """
      INSERT INTO ZCVERS
        (owner_user_id, batch_id, COMPONENT, `RELEASE`, EXTRELEASE, COMP_TYPE)
      VALUES (%s,%s,%s,%s,%s,%s)
    """
    for r in rows:
        cur.execute(sql, (
            owner_user_id, batch_id,
            r.get("COMPONENT") or "", r.get("RELEASE") or "", r.get("EXTRELEASE") or "",
            (r.get("COMP_TYPE") or "")[:1] or None,
        ))

def _call_openwebui_security_analysis(prdvers_rows: list, cvers_rows: list) -> dict:
    if not OPENWEBUI_API_KEY:
        raise HTTPException(500, "OPENWEBUI_API_KEY 未設定")

    lines = ["=== PRDVERS（產品版本） ==="]
    for r in prdvers_rows:
        lines.append(f"ID={r.get('ID')} NAME={r.get('NAME')} VERSION={r.get('VERSION')} "
                      f"VENDOR={r.get('VENDOR')} DESCRIPT={r.get('DESCRIPT')}")
    lines.append("=== CVERS（軟體元件版本） ===")
    for r in cvers_rows:
        lines.append(f"COMPONENT={r.get('COMPONENT')} RELEASE={r.get('RELEASE')} EXTRELEASE={r.get('EXTRELEASE')}")
    user_content = "\n".join(lines)

    payload = {
        "model": OPENWEBUI_MODEL,
        "messages": [
            {"role": "system", "content": SECURITY_SCAN_SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        "files": [{"type": "collection", "id": OPENWEBUI_KB_ID}],
    }
    try:
        r = httpx.post(
            f"{OPENWEBUI_BASE_URL}/api/chat/completions",
            json=payload,
            headers={"Authorization": f"Bearer {OPENWEBUI_API_KEY}"},
            timeout=90,
        )
        r.raise_for_status()
    except httpx.HTTPError as e:
        raise HTTPException(502, f"呼叫 AI 服務失敗：{e}")

    data = r.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError):
        raise HTTPException(502, "AI 回應格式異常")

    text = content.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
    try:
        parsed = json.loads(text)
        summary = parsed.get("summary", "")
        findings = parsed.get("findings", [])
    except Exception:
        # AI 沒有照格式回答時，退化成把整段文字當摘要，避免整個功能直接失敗
        summary = content
        findings = []

    return {"summary": summary, "findings": findings}

@app.post("/api/users/{userid}/security-scan")
def run_security_scan(userid: str):
    prdvers_rows = _rfc_read_table(userid, "PRDVERS", PRDVERS_FIELDS, 500)["rows"]
    cvers_rows   = _rfc_read_table(userid, "CVERS",   CVERS_FIELDS,   500)["rows"]

    analysis = _call_openwebui_security_analysis(prdvers_rows, cvers_rows)

    batch_id = str(uuid.uuid4())
    source_snapshot = json.dumps({"PRDVERS": prdvers_rows, "CVERS": cvers_rows}, ensure_ascii=False)

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT id FROM USRINFO WHERE userid=%s LIMIT 1", (userid,))
        u = cur.fetchone()
        if not u:
            raise HTTPException(404, "user not found")
        owner_user_id = u["id"]

        try:
            conn.begin()
            _insert_prdvers_batch(cur, owner_user_id, batch_id, prdvers_rows)
            _insert_cvers_batch(cur, owner_user_id, batch_id, cvers_rows)
            cur.execute("""
                INSERT INTO security_scans (user_id, batch_id, summary, findings, source_snapshot, model_used)
                VALUES (%s,%s,%s,%s,%s,%s)
            """, (
                owner_user_id, batch_id, analysis["summary"],
                json.dumps(analysis["findings"], ensure_ascii=False),
                source_snapshot, OPENWEBUI_MODEL
            ))
            scan_id = cur.lastrowid
            conn.commit()
        except:
            conn.rollback()
            raise

    return {
        "scan_id": scan_id,
        "batch_id": batch_id,
        "summary": analysis["summary"],
        "findings": analysis["findings"],
        "raw_data": {"PRDVERS": prdvers_rows, "CVERS": cvers_rows},
    }

@app.get("/api/users/{userid}/security-scan/{scan_id}")
def get_security_scan(userid: str, scan_id: int):
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("""
            SELECT s.id, s.batch_id, s.summary, s.findings, s.source_snapshot, s.model_used, s.created_at
              FROM security_scans s
              JOIN USRINFO u ON u.id = s.user_id
             WHERE s.id=%s AND u.userid=%s LIMIT 1
        """, (scan_id, userid))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, "scan not found")
        row["findings"] = json.loads(row["findings"]) if row["findings"] else []
        row["raw_data"] = json.loads(row.pop("source_snapshot")) if row.get("source_snapshot") else {"PRDVERS": [], "CVERS": []}
        return row

@app.get("/api/users/{userid}/security-scan/{scan_id}/export")
def export_security_scan(userid: str, scan_id: int):
    from docx import Document

    scan = get_security_scan(userid, scan_id)

    doc = Document()
    doc.add_heading("SAP 資安漏洞檢測報告", level=1)
    doc.add_paragraph(f"帳號：{userid}")
    doc.add_paragraph(f"檢測時間：{scan['created_at']}")
    doc.add_heading("摘要", level=2)
    doc.add_paragraph(scan["summary"] or "")
    doc.add_heading("發現項目", level=2)
    sev_map = {"high": "高風險", "medium": "中風險", "low": "低風險"}
    if not scan["findings"]:
        doc.add_paragraph("未發現需留意的項目。")
    for f in scan["findings"]:
        p = doc.add_paragraph()
        p.add_run(f"[{sev_map.get(f.get('severity'), f.get('severity',''))}] {f.get('title','')}").bold = True
        doc.add_paragraph(f.get("detail", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"security_scan_{userid}_{scan_id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ==================== Points APIs（統一到 /api/users/{userid}/points...） ====================
# 說明：
# - 正式路徑統一使用 /api/users/{userid}/points...
# - 同時保留舊路徑 /api/points/{userid}/... 作為相容別名（deprecated）。
# - 內部商業邏輯沿用：_coerce_amount、_get_user_id_and_balance、_insert_ledger，
#   以及既有的 credit/debit/refund/set-balance 與 ledger 查詢。

# ---- 型別與小工具 -------------------------------------------------------------

class PointsChange(BaseModel):
    amount: Optional[int] = None
    points: Optional[int] = None
    value:  Optional[int] = None
    delta:  Optional[int] = None
    note: Optional[str] = None

class PointsSetBalance(BaseModel):
    target_balance: Optional[int] = None
    target: Optional[int] = None
    balance: Optional[int] = None
    note: Optional[str] = None

def _coerce_amount(body: PointsChange) -> int:
    v = body.amount or body.points or body.value or body.delta
    if v is None or int(v) <= 0:
        raise HTTPException(400, "amount must be positive integer")
    return int(v)

def _get_user_id_and_balance(cur, userid: str):
    cur.execute("SELECT id, points_balance FROM USRINFO WHERE userid=%s LIMIT 1", (userid,))
    row = cur.fetchone()
    if not row:
        raise HTTPException(404, "user not found")
    return row["id"], int(row["points_balance"])

def _insert_ledger(cur, user_id: int, kind: str, amount: int, balance_after: int, note: Optional[str]):
    cur.execute(
        "INSERT INTO points_ledger (user_id, kind, amount, balance_after, note) VALUES (%s,%s,%s,%s,%s)",
        (user_id, kind, amount, balance_after, note or "")
    )

def _shape_balance_payload(row: dict) -> dict:
    """把 balance 欄位做出兩組 key：balance/updated_at 與 points_balance/points_updated_at。"""
    return {
        "userid":            row["userid"],
        "role":              row["role"],
        "is_active":         row["is_active"],
        "points_balance":    row["points_balance"],
        "points_updated_at": row["points_updated_at"],
        # 新增更直覺的名稱，前端可讀：
        "balance":           row["points_balance"],
        "updated_at":        row["points_updated_at"],
    }

# ---- DB 查詢：餘額與明細（核心查詢，不掛路由） -------------------------------

def _db_get_points_balance(userid: str) -> dict:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("""
            SELECT userid, role, points_balance, points_updated_at, is_active
              FROM USRINFO
             WHERE userid=%s
             LIMIT 1
        """, (userid,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, "user not found")
        return row

def _db_get_points_ledger(
    userid: str,
    kind: Optional[Literal["CREDIT","DEBIT","REFUND"]] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    q: Optional[str] = None,
    page: int = 1,
    page_size: int = 50,
):
    page = max(1, page)
    page_size = max(1, min(200, page_size))
    offset = (page - 1) * page_size

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT id FROM USRINFO WHERE userid=%s", (userid,))
        u = cur.fetchone()
        if not u:
            raise HTTPException(404, "user not found")
        uid = u["id"]

        where = ["user_id=%s"]
        params: list = [uid]
        if kind:
            where.append("kind=%s"); params.append(kind)
        if date_from:
            where.append("created_at >= %s"); params.append(f"{date_from} 00:00:00")
        if date_to:
            where.append("created_at <= %s"); params.append(f"{date_to} 23:59:59")
        if q:
            where.append("note LIKE %s"); params.append(f"%{q}%")
        where_sql = " AND ".join(where)

        cur.execute(f"SELECT COUNT(*) AS c FROM points_ledger WHERE {where_sql}", params)
        total = cur.fetchone()["c"]

        cur.execute(
            f"""SELECT id, kind, amount, balance_after, note, created_at
                  FROM points_ledger
                 WHERE {where_sql}
              ORDER BY created_at DESC, id DESC
                 LIMIT %s OFFSET %s""",
            params + [page_size, offset]
        )
        items = cur.fetchall()

    return {"total": total, "page": page, "page_size": page_size, "items": items}

# ---- 寫入：加點/扣點/退款/設定餘額（核心邏輯，不掛路由） ----------------------

def _op_credit(userid: str, body: PointsChange) -> dict:
    amt = _coerce_amount(body)
    with get_conn() as conn, conn.cursor() as cur:
        try:
            conn.begin()
            user_id, _ = _get_user_id_and_balance(cur, userid)
            cur.execute("""
                UPDATE USRINFO
                   SET points_balance = points_balance + %s,
                       points_updated_at = NOW()
                 WHERE id=%s
            """, (amt, user_id))
            cur.execute("SELECT points_balance FROM USRINFO WHERE id=%s", (user_id,))
            bal = cur.fetchone()["points_balance"]
            _insert_ledger(cur, user_id, "CREDIT", amt, bal, body.note)
            conn.commit()
            return {"ok": True, "points_balance": bal}
        except:
            conn.rollback()
            raise

def _op_debit(userid: str, body: PointsChange) -> dict:
    amt = _coerce_amount(body)
    with get_conn() as conn, conn.cursor() as cur:
        try:
            conn.begin()
            user_id, cur_bal = _get_user_id_and_balance(cur, userid)
            if cur_bal < amt:
                conn.rollback()
                raise HTTPException(402, "Insufficient points")
            cur.execute("""
                UPDATE USRINFO
                   SET points_balance = points_balance - %s,
                       points_updated_at = NOW()
                 WHERE id=%s
            """, (amt, user_id))
            cur.execute("SELECT points_balance FROM USRINFO WHERE id=%s", (user_id,))
            bal = cur.fetchone()["points_balance"]
            _insert_ledger(cur, user_id, "DEBIT", amt, bal, body.note)
            conn.commit()
            return {"ok": True, "points_balance": bal}
        except:
            conn.rollback()
            raise

def _op_refund(userid: str, body: PointsChange) -> dict:
    amt = _coerce_amount(body)
    with get_conn() as conn, conn.cursor() as cur:
        try:
            conn.begin()
            user_id, _ = _get_user_id_and_balance(cur, userid)
            cur.execute("""
                UPDATE USRINFO
                   SET points_balance = points_balance + %s,
                       points_updated_at = NOW()
                 WHERE id=%s
            """, (amt, user_id))
            cur.execute("SELECT points_balance FROM USRINFO WHERE id=%s", (user_id,))
            bal = cur.fetchone()["points_balance"]
            _insert_ledger(cur, user_id, "REFUND", amt, bal, body.note)
            conn.commit()
            return {"ok": True, "points_balance": bal}
        except:
            conn.rollback()
            raise

def _op_set_balance(userid: str, body: PointsSetBalance) -> dict:
    target = (
        body.target_balance if body.target_balance is not None
        else (body.target if body.target is not None else body.balance)
    )
    if target is None or int(target) < 0:
        raise HTTPException(400, "target_balance must be >= 0")
    target = int(target)

    with get_conn() as conn, conn.cursor() as cur:
        try:
            conn.begin()
            user_id, current = _get_user_id_and_balance(cur, userid)
            delta = target - current
            if delta == 0:
                conn.commit()
                return {"ok": True, "points_balance": current, "note": "no change"}

            if delta > 0:
                cur.execute("""
                    UPDATE USRINFO
                       SET points_balance = points_balance + %s,
                           points_updated_at = NOW()
                     WHERE id=%s
                """, (delta, user_id))
                new_balance = current + delta
                _insert_ledger(cur, user_id, "CREDIT", delta, new_balance, body.note or "admin set balance")
            else:
                need = abs(delta)
                if current < need:
                    conn.rollback()
                    raise HTTPException(402, "Insufficient points for set-balance")
                cur.execute("""
                    UPDATE USRINFO
                       SET points_balance = points_balance - %s,
                           points_updated_at = NOW()
                     WHERE id=%s
                """, (need, user_id))
                new_balance = current - need
                _insert_ledger(cur, user_id, "DEBIT", need, new_balance, body.note or "admin set balance")

            conn.commit()
            return {"ok": True, "points_balance": new_balance}
        except:
            conn.rollback()
            raise

# ---- 正式路徑（統一） /api/users/{userid}/points... ----------------------------

@app.get("/api/users/{userid}/points")
def users_points_balance(userid: str):
    row = _db_get_points_balance(userid)
    return _shape_balance_payload(row)

@app.get("/api/users/{userid}/points/history")
def users_points_history(userid: str, limit: int = 100):
    """
    簡化回傳：{ items: [...], total: N }
    內部仍舊使用分頁查詢，這裡固定 page=1, page_size=limit。
    """
    data = _db_get_points_ledger(userid, None, None, None, None, page=1, page_size=limit)
    return {"items": data.get("items", []), "total": data.get("total", 0)}

@app.post("/api/users/{userid}/points/adjust")
def users_points_adjust(userid: str, body: dict = Body(...)):
    """
    body: { amount, kind: 'CREDIT'|'DEBIT'|'REFUND', note? }
    """
    kind = (body.get("kind") or "").upper()
    payload = PointsChange(
        amount=body.get("amount"),
        note=body.get("note"),
        points=body.get("points"),
        value=body.get("value"),
        delta=body.get("delta"),
    )
    if kind == "CREDIT":
        return _op_credit(userid, payload)
    if kind == "DEBIT":
        return _op_debit(userid, payload)
    if kind == "REFUND":
        return _op_refund(userid, payload)
    raise HTTPException(400, "kind must be one of: CREDIT|DEBIT|REFUND")

# 明確動作路徑（保留，方便直接呼叫特定操作）
@app.post("/api/users/{userid}/points/grant")
def users_points_grant(userid: str, body: PointsChange = Body(...)):
    return _op_credit(userid, body)

@app.post("/api/users/{userid}/points/debit")
def users_points_debit(userid: str, body: PointsChange = Body(...)):
    return _op_debit(userid, body)

@app.post("/api/users/{userid}/points/refund")
def users_points_refund(userid: str, body: PointsChange = Body(...)):
    return _op_refund(userid, body)

@app.post("/api/users/{userid}/points/set-balance")
def users_points_set_balance(userid: str, body: PointsSetBalance = Body(...)):
    return _op_set_balance(userid, body)

# ---- 舊路徑（相容別名；建議逐步淘汰） /api/points/{userid}/... --------------
# 這裡直接委派到正式路徑實作；回傳會比舊版多出 balance/updated_at 等欄位（相容擴充）。

@app.get("/api/points/{userid}/balance")
def legacy_points_balance(userid: str):
    return users_points_balance(userid)

@app.get("/api/points/{userid}/ledger")
def legacy_points_ledger(
    userid: str,
    kind: Optional[Literal["CREDIT","DEBIT","REFUND"]] = Query(default=None),
    date_from: Optional[str] = Query(default=None),
    date_to: Optional[str] = Query(default=None),
    q: Optional[str] = Query(default=None),
    page: int = 1,
    page_size: int = 50,
):
    # 保留舊參數語意；直接用核心查詢回傳完整分頁結構
    return _db_get_points_ledger(userid, kind, date_from, date_to, q, page, page_size)

@app.post("/api/points/{userid}/credit")
def legacy_points_credit(userid: str, body: PointsChange):
    return _op_credit(userid, body)

@app.post("/api/points/{userid}/debit")
def legacy_points_debit(userid: str, body: PointsChange):
    return _op_debit(userid, body)

@app.post("/api/points/{userid}/refund")
def legacy_points_refund(userid: str, body: PointsChange):
    return _op_refund(userid, body)

@app.post("/api/points/{userid}/set-balance")
def legacy_points_set_balance(userid: str, body: PointsSetBalance):
    return _op_set_balance(userid, body)

@app.post("/api/points/{userid}/adjust")
def legacy_points_adjust(userid: str, payload: dict = Body(...)):
    # 舊版可能用 op/type 指定動作；這裡轉為新 adjust 規格
    op = (payload.get("op") or payload.get("type") or "").lower()
    kind = "CREDIT" if op in ("credit","add","grant","plus") else \
           "DEBIT"  if op in ("debit","sub","charge","minus") else \
           "REFUND" if op in ("refund","rollback","back") else None
    if not kind:
        raise HTTPException(400, "op must be one of: credit|debit|refund")
    return users_points_adjust(userid, {
        "amount": payload.get("amount") or payload.get("points") or payload.get("value") or payload.get("delta"),
        "note":   payload.get("note"),
        "kind":   kind,
    })